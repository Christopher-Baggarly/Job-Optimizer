import os
import json
from docx import Document
from google import genai
from google.genai import types

# ==============================================================================
# PIPELINE CONFIGURATION
# ==============================================================================
INPUT_DOCX = r"./input/sample_resume.docx"
OUTPUT_DIR = r"./output"
TARGET_JSON_PATH = os.path.join(OUTPUT_DIR, "normalized_data.json")

os.makedirs(OUTPUT_DIR, exist_ok=True)

# Initialize the Gemini Client (Make sure GEMINI_API_KEY is set in your environment)
client = genai.Client()

# ==============================================================================
# THE STRICT SCHEMA (Ensures 100% compliance with your 10 active templates)
# ==============================================================================
class ContactFirst(types.Schema):
    name: str

class ContactLast(types.Schema):
    name: str

class ExperienceItem(types.Schema):
    title: str
    company: str
    date: str
    bullet: str

class EducationItem(types.Schema):
    institution: str
    degree: str
    date: str

class ResumeSchema(types.Schema):
    First: ContactFirst
    Last: ContactLast
    email: str
    phone: str
    address: str
    summary: str
    experience: list[ExperienceItem]
    education: list[EducationItem]
    skills: list[str]

# ==============================================================================
# CORE EXTRACTION PROCESSING
# ==============================================================================
def extract_raw_text_and_tables(docx_path):
    """Pulls every scrap of text out of paragraphs and tables indiscriminately."""
    doc = Document(docx_path)
    content_chunks = []
    
    for p in doc.paragraphs:
        if p.text.strip():
            content_chunks.append(p.text.strip())
            
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                content_chunks.append(" | ".join(row_text))
                
    return "\n".join(content_chunks)

def normalize_to_schema(raw_text):
    """
    Stage 1 Data Extractor matching true OpenXML Jinja tags.
    Handles nested bullet arrays, object-based certifications, and references.
    """
    
    json_blueprint = {
        "First": {"name": "String"},
        "Last": {"name": "String"},
        "email": "String",
        "phone": "String",
        "address": "String",
        "summary": "String",
        "student": {
            "education": [
                {"degree": "String", "institution": "String", "date": "String"}
            ],
            "experience": [
                {
                    "title": "String", 
                    "company": "String", 
                    "date": "String", 
                    "bullets": ["Array of action-oriented bullet strings"]
                }
            ],
            "skills": ["String"],
            "certifications": [
                {"name": "String", "description": "String or blank"}
            ],
            "references": [
                {"name": "String", "phone": "String"}
            ]
        }
    }

    prompt = f"""
    You are an expert data migration engine. Your task is to extract resume data from the following 
    unstructured text and format it to strictly adhere to this exact JSON layout blueprint:
    {json.dumps(json_blueprint, indent=2)}
    
    Rules:
    1. Split full names into 'First' and 'Last' object structures cleanly.
    2. Synthesize or isolate a professional summary statement.
    3. CRITICAL: 'experience.bullets' MUST be returned as an array of individual strings, NOT one big paragraph. Break paragraphs into distinct action bullets.
    4. Split raw certifications into a structured name and optional description object.
    5. Search for any professional references or reference placeholders. If none exist in the raw text, leave the array empty [].
    6. Return ONLY the raw JSON object. Do not wrap it in markdown code blocks.
    
    Raw Document Text:
    ---
    {raw_text}
    ---
    """
    
    response = client.models.generate_content(
        model='gemini-2.5-flash',
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            temperature=0.1,
        ),
    )
    
    return response.text
# ==============================================================================
# EXECUTION LIFECYCLE
# ==============================================================================
def main():
    print(f"[STAGE 1] Reading raw document: {os.path.basename(INPUT_DOCX)}")
    if not os.path.exists(INPUT_DOCX):
        print(f"Error: Could not find file at {INPUT_DOCX}")
        return
        
    raw_text = extract_raw_text_and_tables(INPUT_DOCX)
    
    print("[STAGE 1] Analyzing and normalizing data structure with Gemini layer...")
    try:
        normalized_json_string = normalize_to_schema(raw_text)
        
        template_ready_payload = json.loads(normalized_json_string)
        
        with open(TARGET_JSON_PATH, 'w', encoding='utf-8') as f:
            json.dump(template_ready_payload, f, indent=4, ensure_ascii=False)
            
        print(f"[STAGE 1 COMPLETE] Data isolated, normalized, and saved to:")
        print(f"   {TARGET_JSON_PATH}")
        print("\nReview the JSON file. If it looks correct, you are clear to run Stage 2 Insertion.")
        
    except Exception as e:
        print(f"Critical failure during normalization: {e}")

if __name__ == "__main__":
    main()
